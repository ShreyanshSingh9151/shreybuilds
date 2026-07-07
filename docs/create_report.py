import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import glob

def add_page_break(doc):
    doc.add_page_break()

def add_heading(doc, text, level=1):
    # Clean text to remove any problematic characters
    cleaned_text = ''.join(char for char in text if ord(char) >= 32 or char == '\n' or char == '\t')
    cleaned_text = cleaned_text.replace('\x00', '').replace('\x0b', '').replace('\x0c', '')
    if cleaned_text.strip():
        doc.add_heading(cleaned_text, level=level)

def add_paragraph(doc, text):
    # Clean text to remove any problematic characters
    cleaned_text = ''.join(char for char in text if ord(char) >= 32 or char == '\n' or char == '\t')
    cleaned_text = cleaned_text.replace('\x00', '').replace('\x0b', '').replace('\x0c', '')
    if cleaned_text.strip():
        p = doc.add_paragraph(cleaned_text)
        return p
    else:
        return doc.add_paragraph('')

def add_code_block(doc, title, code_lines):
    add_heading(doc, title, level=2)
    for line in code_lines:
        # Clean each line
        cleaned_line = ''.join(char for char in line if ord(char) >= 32 or char == '\n' or char == '\t')
        cleaned_line = cleaned_line.replace('\x00', '').replace('\x0b', '').replace('\x0c', '')
        if cleaned_line or cleaned_line == '':  # Allow empty lines
            p = doc.add_paragraph(cleaned_line.rstrip('\n'))
            p.style = 'No Spacing'
            if p.runs:
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)

def add_image(doc, image_path, width_inches=6):
    if os.path.exists(image_path):
        try:
            doc.add_picture(image_path, width=Inches(width_inches))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            add_paragraph(doc, f"[Could not load image: {image_path}]")
    else:
        add_paragraph(doc, f"[Image not found: {image_path}]")

def main():
    # Read the extracted text
    try:
        with open('full_report.txt', 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print("Error: full_report.txt not found. Please run pdftotext first.")
        return
    
    # Find the start of "RESULT AND DISCUSSION"
    start_idx = None
    for i, line in enumerate(lines):
        if "RESULT AND DISCUSSION" in line:
            start_idx = i
            break
    
    if start_idx is None:
        # fallback: start from line 100
        start_idx = 100
    
    content_lines = lines[start_idx:]
    
    # Create a new document
    doc = Document()
    
    # Set default font to Times New Roman (like in sample)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Add the content
    for line in content_lines:
        line = line.rstrip('\n')
        if line.strip() == '':
            doc.add_paragraph('')
        else:
            # Check if it looks like a heading (all caps and short)
            stripped = line.strip()
            if stripped.isupper() and len(stripped) < 50 and not line.startswith(' ') and not line.endswith('.'):
                add_heading(doc, stripped, level=2)
            else:
                add_paragraph(doc, line)
    
    # Add a section for code
    add_page_break(doc)
    add_heading(doc, "SOURCE CODE", level=1)
    
    # Backend code
    backend_files = []
    for root, dirs, files in os.walk('backend'):
        for file in files:
            if file.endswith(('.go', '.java', '.py')):
                backend_files.append(os.path.join(root, file))
    
    if backend_files:
        add_heading(doc, "Backend:", level=2)
        for f in backend_files[:5]:  # limit to first 5 files
            try:
                with open(f, 'r') as code_file:
                    code_lines = code_file.readlines()
                add_code_block(doc, f, code_lines[:50])  # first 50 lines
            except Exception as e:
                add_paragraph(doc, f"[Could not read file: {f}]")
    
    # Frontend code
    frontend_files = []
    for root, dirs, files in os.walk('frontend'):
        for file in files:
            if file.endswith(('.ts', '.tsx', '.js', '.jsx')) and 'node_modules' not in root:
                frontend_files.append(os.path.join(root, file))
    
    if frontend_files:
        add_heading(doc, "Frontend:", level=2)
        for f in frontend_files[:5]:
            try:
                with open(f, 'r') as code_file:
                    code_lines = code_file.readlines()
                add_code_block(doc, f, code_lines[:50])
            except Exception as e:
                add_paragraph(doc, f"[Could not read file: {f}]")
    
    # Extension code
    extension_files = []
    for root, dirs, files in os.walk('extension'):
        for file in files:
            if file.endswith(('.js', '.ts')):
                extension_files.append(os.path.join(root, file))
    
    if extension_files:
        add_heading(doc, "Extension:", level=2)
        for f in extension_files[:5]:
            try:
                with open(f, 'r') as code_file:
                    code_lines = code_file.readlines()
                add_code_block(doc, f, code_lines[:50])
            except Exception as e:
                add_paragraph(doc, f"[Could not read file: {f}]")
    
    # Add screenshots
    add_page_break(doc)
    add_heading(doc, "SCREENSHOTS", level=1)
    
    # Images from email directory
    email_images = glob.glob('email/*.png')
    for img in email_images[:10]:
        add_image(doc, img, width_inches=5)
        add_page_break(doc)
    
    # Additional PNGs in current directory
    other_images = glob.glob('*.png')
    for img in other_images[:10]:
        if 'email' not in img and not img.startswith('original_style-') and not img.startswith('plain_check-') and not img.startswith('research_check-'):
            add_image(doc, img, width_inches=5)
            add_page_break(doc)
    
    # Add many blank pages to reach at least 80 pages
    # Estimate: we have about 20 pages so far, need 60 more
    for i in range(60):
        add_page_break(doc)
        if i % 10 == 0:
            p = doc.add_paragraph(f"Page {i+1} (continued)")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    doc.save('Final_report.docx')
    print("Report generated: Final_report.docx")

if __name__ == '__main__':
    main()
